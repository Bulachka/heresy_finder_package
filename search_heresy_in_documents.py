import os
import re
from collections import defaultdict
from docx import Document
from docx.shared import Pt
import fitz  # PyMuPDF for handling PDFs
from patterns import raw_patterns

# === CONFIGURATION ===
folder_path = "input_txt"  # <- Change to your folder path if needed
output_file = os.path.join("output", "search_results.docx")
os.makedirs(os.path.dirname(output_file), exist_ok=True)

# === COMPILE REGEX PATTERNS ===
regex_patterns = [(pattern, re.compile(pattern, re.IGNORECASE)) for pattern in raw_patterns]

# === CREATE WORD DOCUMENT ===
doc = Document()
doc.add_heading("Search Results", level=1)

# === DICTIONARY TO COUNT MATCHES ===
match_counts = defaultdict(int)

# === FUNCTION TO HIGHLIGHT AND WRITE MATCH IN WORD ===
def add_match_to_doc(filename, context_text, match_text):
    doc.add_heading(f"File: {filename}", level=2)
    paragraph = doc.add_paragraph()
    try:
        before, after = context_text.split("**" + match_text + "**")
    except ValueError:
        before = context_text
        after = ""
        match_text = ""

    run_before = paragraph.add_run(before)
    run_match = paragraph.add_run(match_text)
    run_match.bold = True
    run_after = paragraph.add_run(after)

    for run in [run_before, run_match, run_after]:
        run.font.size = Pt(11)

# === FUNCTION TO PROCESS TEXT ===
def process_text(filename, text):
    for pattern_str, regex in regex_patterns:
        for match in regex.finditer(text):
            start = max(0, match.start() - 1000)
            end = min(len(text), match.end() + 1000)
            context = text[start:match.start()] + "**" + match.group(0) + "**" + text[match.end():end]
            add_match_to_doc(filename, context, match.group(0))
            match_counts[pattern_str] += 1

# === MAIN PROCESSING LOOP ===
for filename in os.listdir(folder_path or "."):
    file_path = os.path.join(folder_path, filename)

    if filename.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            process_text(filename, text)

    elif filename.endswith(".pdf"):
        doc_pdf = fitz.open(file_path)
        text = ""
        for page in doc_pdf:
            text += page.get_text()
        process_text(filename, text)

# === ADD MATCH COUNT TABLE TO DOCUMENT ===
doc.add_page_break()
doc.add_heading("Match Count Summary", level=1)

table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Pattern'
hdr_cells[1].text = 'Match Count'

for pattern, count in match_counts.items():
    row_cells = table.add_row().cells
    row_cells[0].text = pattern
    row_cells[1].text = str(count)

# === SAVE THE WORD DOCUMENT ===
doc.save(output_file)
print(f"✅ Results saved to: {output_file}")
