import sys
import re
from docx import Document
from docx.shared import RGBColor
from collections import defaultdict

# Heretical word patterns (all should match from the beginning of a word)
raw_patterns = [
    "chetzer.*", "kætzer.*", "kätzer.*", "keczer.*", "keczir.*", "keczzer.*",
    "ketzcer.*", "ketzcir.*", "ketzeer.*", "ketzer.*", "khetzer.*", "ketczer.*",
    "haeres\\w+", "haeret.*", "heret.*", "heres\\w+",
    "huss.*"
]

# Compile pattern for full-word match
regex = re.compile(r'^(' + '|'.join(raw_patterns) + r')$', re.IGNORECASE)
pattern_stats = defaultdict(int)

if len(sys.argv) != 3:
    print("Usage: python highlight_heresy.py input.docx output.docx")
    sys.exit(1)

input_file = sys.argv[1]
output_file = sys.argv[2]

# Load the input DOCX
doc = Document(input_file)
new_doc = Document()

for para in doc.paragraphs:
    new_para = new_doc.add_paragraph()
    
    for run in para.runs:
        # Split text in the run into words + punctuation
        words = re.findall(r'\w+|\W+', run.text)

        for word in words:
            r = new_para.add_run(word)

            # Preserve formatting
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
            r.font.name = run.font.name
            r.font.size = run.font.size

            # Preserve color if it's not already red
            if run.font.color and run.font.color.rgb != RGBColor(255, 0, 0):
                r.font.color.rgb = run.font.color.rgb

            # Match pattern and highlight red if it matches
            if regex.match(word.strip()):
                r.font.color.rgb = RGBColor(255, 0, 0)

                for pattern in raw_patterns:
                    if re.fullmatch(pattern, word.strip(), re.IGNORECASE):
                        pattern_stats[pattern] += 1
                        break

# Add match stats as a table at the end
new_doc.add_paragraph("\nMatch counts by pattern:")
table = new_doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = 'Pattern'
table.rows[0].cells[1].text = 'Count'

for pattern in raw_patterns:
    row = table.add_row().cells
    row[0].text = pattern
    row[1].text = str(pattern_stats[pattern])

# Save the output
new_doc.save(output_file)
print(f"✅ Done! Highlighted file saved to {output_file}")
